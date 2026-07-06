"""build_bilingual_py.py
python-docx template for generating bilingual or clean translated Word documents:
  - EN-AR file: left column English, right column Arabic
  - ZH-AR file: left column Chinese, right column Arabic
  - Clean file: target-language translation only, no source-text column

All body paragraphs use 1.2x line spacing (see LINE_SPACING_MULTIPLE below).

Usage:
  1. Copy this file to your working directory.
  2. Edit the DOCUMENTS list below with your content.
  3. Install python-docx into the active managed Python if needed:
       python -m pip install python-docx
  4. Run:
       python build_bilingual_py.py

Output: one .docx per entry in DOCUMENTS, written to the current directory.
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# 1.2x line spacing, applied to every paragraph in the document.
LINE_SPACING_MULTIPLE = 1.2

# =============================================================================
# CONTENT - EDIT THIS SECTION
# =============================================================================
# Each document is one output file.
# layout="side-by-side" uses rows as aligned [left_text, right_text] pairs.
# layout="clean" uses paragraphs as target-language text only. If paragraphs is
# omitted, clean output is derived from rows: left text for EN/ZH, right text for AR.
# `mode` controls the pair: "EN-AR" (English left / Arabic right)
#                          "ZH-AR" (Chinese left / Arabic right)
# `target_language` for clean output: "EN", "ZH", or "AR".

DOCUMENTS = [
    {
        "layout": "side-by-side",
        "mode": "EN-AR",
        "basename": "sample_EN-AR",
        "title": "Bilingual Document - English / Arabic",
        "subtitle": "Working translation - not certified",
        "confidentiality": "DRAFT - CONFIDENTIAL",
        "rows": [
            ["Decision No. DE-1447-00053", "قرار رقم DE-1447-00053"],
            ["The General Authority for Competition, after reviewing the file,", "اطلعت الهيئة العامة للمنافسة على الملف،"],
            ["decides the following:", "تقرر ما يلي:"],
            ["First: The respondent shall cease the conduct.", "أولاً: يلتزم المخالف بإيقاف السلوك."],
            ["Second: A fine of SAR 100,000 is imposed.", "ثانياً: تُفرض غرامة قدرها 100,000 ريال سعودي."],
        ],
        "translators_note": [
            "Translator's note (for counsel verification - this is a working translation, not a certified/sworn translation):",
            "1. Judgment calls: none in this sample.",
            "2. Ambiguities / source defects: none observed.",
            "3. Date conversions: none in this sample.",
            "4. Terminology choices: 'General Authority for Competition' used throughout.",
            "5. Backend used: python-docx.",
        ],
    },
    {
        "layout": "side-by-side",
        "mode": "ZH-AR",
        "basename": "sample_ZH-AR",
        "title": "对照文档 - 中文 / 阿拉伯语",
        "subtitle": "工作译本 - 非公证翻译",
        "confidentiality": "草案 - 机密",
        "rows": [
            ["第 DE-1447-00053 号决定", "قرار رقم DE-1447-00053"],
            ["竞争总局在审阅案卷后，", "اطلعت الهيئة العامة للمنافسة على الملف،"],
            ["决定如下：", "تقرر ما يلي:"],
            ["一、被调查方应停止上述行为。", "أولاً: يلتزم المخالف بإيقاف السلوك."],
            ["二、处以 100,000 沙特里亚尔罚款。", "ثانياً: تُفرض غرامة قدرها 100,000 ريال سعودي."],
        ],
        "translators_note": [
            "译注（供律师复核 - 本译本为工作翻译，非公证/宣誓翻译）：",
            "1. 判断取舍：本样例无。",
            "2. 歧义/源文瑕疵：未发现。",
            "3. 日期换算：本样例无。",
            "4. 术语选择：全文统一译法。",
            "5. 使用的后端：python-docx。",
        ],
    },
    {
        "layout": "clean",
        "mode": "EN-AR",
        "target_language": "EN",
        "basename": "sample_EN_clean",
        "title": "English Translation",
        "subtitle": "Clean working translation - not certified",
        "confidentiality": "DRAFT - CONFIDENTIAL",
        "paragraphs": [
            "Decision No. DE-1447-00053",
            "The General Authority for Competition, after reviewing the file,",
            "decides the following:",
            "First: The respondent shall cease the conduct.",
            "Second: A fine of SAR 100,000 is imposed.",
        ],
        "translators_note": [
            "Translator's note (for counsel verification - this is a working translation, not a certified/sworn translation):",
            "1. Judgment calls: none in this sample.",
            "2. Ambiguities / source defects: none observed.",
            "3. Date conversions: none in this sample.",
            "4. Terminology choices: 'General Authority for Competition' used throughout.",
            "5. Backend used: python-docx.",
        ],
    },
]

# =============================================================================
# FONTS AND SIZES
# =============================================================================

FONT_LATIN = "Times New Roman"
FONT_CJK = "SimSun"  # 宋体
FONT_ARABIC = "Amiri"  # install Amiri, or fall back to "Traditional Arabic"
SIZE_BODY = Pt(12)
SIZE_HEADING = Pt(16)
SIZE_SUBTITLE = Pt(11)
SIZE_NOTE = Pt(10)
COLOR_GREY = RGBColor(0x80, 0x80, 0x80)
HEADER_FILL = "F2F2F2"

# =============================================================================
# SPACING HELPER
# =============================================================================


def set_1_2_spacing(paragraph):
    """Apply 1.2x multiple line spacing to a paragraph."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING_MULTIPLE


# =============================================================================
# RTL HELPERS (python-docx has no high-level RTL API)
# =============================================================================


def set_rtl_paragraph(paragraph):
    """Mark a paragraph as right-to-left (w:bidi)."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def set_rtl_run(run):
    """Mark a run as right-to-left (w:rtl)."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def set_complex_script_font(run, font_name):
    """Set complex-script (cs), ascii, and hAnsi font for a run."""
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cjk_font(run, latin_name, cjk_name):
    """Set latin and eastAsia font for a run containing CJK characters."""
    run.font.name = latin_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), latin_name)
    rFonts.set(qn("w:hAnsi"), latin_name)
    rFonts.set(qn("w:eastAsia"), cjk_name)


def set_cell_shading(cell, fill_hex):
    """Apply background shading to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_width(cell, pct):
    """Set table cell width as a percentage (0-100)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(pct * 50)))  # pct*50 gives the fiftieths-of-a-percent unit
    tcW.set(qn("w:type"), "pct")
    tcPr.append(tcW)


def add_arabic_paragraph(cell_or_doc, text, bold=False, size=SIZE_BODY):
    """Add an RTL Arabic paragraph to a cell or document body."""
    p = cell_or_doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    set_rtl_paragraph(p)
    set_1_2_spacing(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    set_complex_script_font(run, FONT_ARABIC)
    set_rtl_run(run)
    return p


def add_latin_paragraph(cell_or_doc, text, bold=False, size=SIZE_BODY):
    """Add an LTR Latin paragraph to a cell or document body."""
    p = cell_or_doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_1_2_spacing(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = FONT_LATIN
    return p


def add_cjk_paragraph(cell_or_doc, text, bold=False, size=SIZE_BODY):
    """Add an LTR Chinese paragraph to a cell or document body."""
    p = cell_or_doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_1_2_spacing(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    set_cjk_font(run, FONT_LATIN, FONT_CJK)
    return p


def target_language(doc_def):
    """Return the clean-output target language: EN, ZH, or AR."""
    if doc_def.get("target_language"):
        return doc_def["target_language"]
    if doc_def.get("mode") == "ZH-AR":
        return "ZH"
    if doc_def.get("mode") == "AR":
        return "AR"
    return "EN"


def is_clean_layout(doc_def):
    """Return True when the document should be rendered as target text only."""
    return doc_def.get("layout") == "clean"


def add_translated_paragraph(doc, item, language):
    """Add one paragraph in the clean translated layout."""
    data = {"text": item} if isinstance(item, str) else item
    text = data.get("text", "")
    bold = data.get("bold", False)
    if language == "AR":
        return add_arabic_paragraph(doc, text, bold=bold)
    if language == "ZH":
        return add_cjk_paragraph(doc, text, bold=bold)
    return add_latin_paragraph(doc, text, bold=bold)


def add_clean_body(doc, doc_def):
    """Add target-language-only body paragraphs."""
    language = target_language(doc_def)
    paragraphs = doc_def.get("paragraphs")
    if paragraphs is None:
        paragraphs = [
            right if language == "AR" else left
            for left, right in doc_def.get("rows", [])
        ]
    for item in paragraphs:
        add_translated_paragraph(doc, item, language)


def add_side_by_side_body(doc, doc_def):
    """Add the two-column bilingual comparison table."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    LEFT_WIDTH = 55
    RIGHT_WIDTH = 45

    # Header row
    hdr_cells = table.rows[0].cells
    left_label = "中文" if doc_def["mode"] == "ZH-AR" else "English"
    # Left header (LTR)
    set_cell_width(hdr_cells[0], LEFT_WIDTH)
    set_cell_shading(hdr_cells[0], HEADER_FILL)
    hdr_p0 = hdr_cells[0].paragraphs[0]
    hdr_p0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_1_2_spacing(hdr_p0)
    hdr_r0 = hdr_p0.add_run(left_label)
    hdr_r0.bold = True
    hdr_r0.font.size = SIZE_BODY
    if doc_def["mode"] == "ZH-AR":
        set_cjk_font(hdr_r0, FONT_LATIN, FONT_CJK)
    else:
        hdr_r0.font.name = FONT_LATIN
    # Right header (RTL Arabic)
    set_cell_width(hdr_cells[1], RIGHT_WIDTH)
    set_cell_shading(hdr_cells[1], HEADER_FILL)
    hdr_p1 = hdr_cells[1].paragraphs[0]
    hdr_p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    set_rtl_paragraph(hdr_p1)
    set_1_2_spacing(hdr_p1)
    hdr_r1 = hdr_p1.add_run("العربية")
    hdr_r1.bold = True
    hdr_r1.font.size = SIZE_BODY
    set_complex_script_font(hdr_r1, FONT_ARABIC)
    set_rtl_run(hdr_r1)

    # Body rows
    for left_text, right_text in doc_def["rows"]:
        row = table.add_row()
        left_cell = row.cells[0]
        right_cell = row.cells[1]
        set_cell_width(left_cell, LEFT_WIDTH)
        set_cell_width(right_cell, RIGHT_WIDTH)
        # Clear default empty paragraph in each cell
        left_cell.paragraphs[0].text = ""
        right_cell.paragraphs[0].text = ""
        # Left content
        if doc_def["mode"] == "ZH-AR":
            add_cjk_paragraph(left_cell, left_text)
        else:
            add_latin_paragraph(left_cell, left_text)
        # Right Arabic content
        add_arabic_paragraph(right_cell, right_text)
        # Remove the leftover empty first paragraph in each cell
        for cell in (left_cell, right_cell):
            if len(cell.paragraphs) > 1 and not cell.paragraphs[0].text:
                cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)


# =============================================================================
# DOCUMENT BUILDER
# =============================================================================


def build_document(doc_def):
    doc = Document()
    language = target_language(doc_def)

    # Page setup: A4 portrait, 1 inch margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Header
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_1_2_spacing(header_p)
    header_run = header_p.add_run(doc_def["confidentiality"])
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = COLOR_GREY
    header_run.font.name = FONT_LATIN

    # Footer with page numbers
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_1_2_spacing(footer_p)
    r1 = footer_p.add_run("Page ")
    r1.font.size = Pt(9)
    r1.font.name = FONT_LATIN
    # PAGE field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_page = footer_p.add_run()
    r_page._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    r_page._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_page._r.append(fld_end)
    r_page.font.size = Pt(9)
    r_page.font.name = FONT_LATIN
    r2 = footer_p.add_run(" of ")
    r2.font.size = Pt(9)
    r2.font.name = FONT_LATIN
    r3 = footer_p.add_run()
    b2 = OxmlElement("w:fldChar")
    b2.set(qn("w:fldCharType"), "begin")
    r3._r.append(b2)
    i2 = OxmlElement("w:instrText")
    i2.text = "NUMPAGES"
    r3._r.append(i2)
    e2 = OxmlElement("w:fldChar")
    e2.set(qn("w:fldCharType"), "end")
    r3._r.append(e2)
    r3.font.size = Pt(9)
    r3.font.name = FONT_LATIN

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_1_2_spacing(title_p)
    title_run = title_p.add_run(doc_def["title"])
    title_run.bold = True
    title_run.font.size = SIZE_HEADING
    if language == "ZH":
        set_cjk_font(title_run, FONT_LATIN, FONT_CJK)
    elif language == "AR":
        set_complex_script_font(title_run, FONT_ARABIC)
        set_rtl_run(title_run)
        set_rtl_paragraph(title_p)
    else:
        title_run.font.name = FONT_LATIN

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_1_2_spacing(sub_p)
    sub_run = sub_p.add_run(doc_def["subtitle"])
    sub_run.italic = True
    sub_run.font.size = SIZE_SUBTITLE
    if language == "ZH":
        set_cjk_font(sub_run, FONT_LATIN, FONT_CJK)
    elif language == "AR":
        set_complex_script_font(sub_run, FONT_ARABIC)
        set_rtl_run(sub_run)
        set_rtl_paragraph(sub_p)
    else:
        sub_run.font.name = FONT_LATIN

    # Spacer
    spacer1 = doc.add_paragraph()
    set_1_2_spacing(spacer1)

    if is_clean_layout(doc_def):
        add_clean_body(doc, doc_def)
    else:
        add_side_by_side_body(doc, doc_def)

    # Spacer
    spacer2 = doc.add_paragraph()
    set_1_2_spacing(spacer2)

    # Translator's note
    if language == "ZH":
        note_heading_text = "译注"
    elif language == "AR":
        note_heading_text = "ملاحظة المترجم"
    else:
        note_heading_text = "Translator's Note"
    nh = doc.add_paragraph()
    set_1_2_spacing(nh)
    if language == "AR":
        nh.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        set_rtl_paragraph(nh)
    nh_run = nh.add_run(note_heading_text)
    nh_run.bold = True
    nh_run.font.size = SIZE_BODY
    if language == "ZH":
        set_cjk_font(nh_run, FONT_LATIN, FONT_CJK)
    elif language == "AR":
        set_complex_script_font(nh_run, FONT_ARABIC)
        set_rtl_run(nh_run)
    else:
        nh_run.font.name = FONT_LATIN

    for line in doc_def["translators_note"]:
        np = doc.add_paragraph()
        set_1_2_spacing(np)
        if language == "AR":
            np.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            set_rtl_paragraph(np)
        # Naive script segmentation: ASCII/Latin/CJK as one run, Arabic as RTL run.
        import re

        arabic_re = re.compile(r"[\u0600-\u06FF\u0750-\u077F]+")
        last = 0
        for m in arabic_re.finditer(line):
            if m.start() > last:
                seg = line[last:m.start()]
                r = np.add_run(seg)
                r.font.size = SIZE_NOTE
                if language == "ZH":
                    set_cjk_font(r, FONT_LATIN, FONT_CJK)
                else:
                    r.font.name = FONT_LATIN
            r_ar = np.add_run(m.group())
            r_ar.font.size = SIZE_NOTE
            set_complex_script_font(r_ar, FONT_ARABIC)
            set_rtl_run(r_ar)
            last = m.end()
        if last < len(line):
            r = np.add_run(line[last:])
            r.font.size = SIZE_NOTE
            if language == "ZH":
                set_cjk_font(r, FONT_LATIN, FONT_CJK)
            else:
                r.font.name = FONT_LATIN

    return doc


def main():
    for doc_def in DOCUMENTS:
        doc = build_document(doc_def)
        out_path = os.path.join(os.getcwd(), f"{doc_def['basename']}.docx")
        doc.save(out_path)
        print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
